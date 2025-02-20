import os
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route("/")
def home():
    return jsonify({"message": "✅ Flask Word 合併 API 已啟動！"})

@app.route("/upload", methods=["POST"])
def upload_files():
    if "files" not in request.files:
        return jsonify({"error": "❌ 未收到檔案"}), 400

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "❌ 未收到檔案"}), 400

    file_paths = []
    for file in files:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)
        file_paths.append(file_path)

    output_path = os.path.join(OUTPUT_FOLDER, "merged_document.docx")
    merge_docs(file_paths, output_path)
    
    return send_file(output_path, as_attachment=True)

def merge_docs(files, output_path):
    merged_doc = Document(files[0])  # 以第一個文件為主文件

    for file in files[1:]:
        merged_doc.add_page_break()  # 插入分頁符號
        doc = Document(file)

        for para in doc.paragraphs:
            new_para = merged_doc.add_paragraph()
            new_para.add_run(para.text).font.size = Pt(12)  # 確保字體大小統一
        
        # 保留表格
        for table in doc.tables:
            new_table = merged_doc.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                new_row = new_table.add_row()
                for cell_idx, cell in enumerate(row.cells):
                    new_row.cells[cell_idx].text = cell.text

    merged_doc.save(output_path)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    print(f"🚀 Flask 伺服器啟動中，監聽 PORT: {port}")
    app.run(host="0.0.0.0", port=port)
